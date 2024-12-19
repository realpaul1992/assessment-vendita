import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
import os
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Mappatura dei valori per la scala 1-5
scala_risposte = {
    1: "Assente/Inesistente",
    2: "Molto poco strutturato",
    3: "Presente ma non ottimizzato",
    4: "Ben strutturato",
    5: "Eccellente/ottimizzato"
}

def selettore_valori_domanda(domanda, key):
    val = st.radio(domanda, options=[1,2,3,4,5], index=2, format_func=lambda x: scala_risposte[x], key=key)
    return val

aree = [
    "Visione e Strategia di Vendita",
    "Processi di Vendita",
    "Script e Protocolli di Vendita",
    "Tecnologia e CRM",
    "Formazione e Sviluppo del Team di Vendita",
    "Misurazione e KPI"
]

domande = {
    "Visione e Strategia di Vendita": {
        "principali": [
            "Esiste una visione chiara e condivisa degli obiettivi di vendita?",
            "I target di vendita sono coerenti con le risorse e comunicati al team?",
            "C'Ã¨ una strategia di vendita chiara per raggiungere gli obiettivi?",
            "Lâ€™efficacia delle strategie viene monitorata regolarmente?",
            "Il management fornisce linee guida e supporto chiari?"
        ],
        "controllo": "Hai una documentazione scritta che descrive obiettivi e strategie?"
    },
    "Processi di Vendita": {
        "principali": [
            "Il percorso del cliente Ã¨ documentato e noto al team?",
            "Esiste un metodo chiaro per qualificare i lead?",
            "Ci sono checklist o step operativi per le fasi di vendita?",
            "Il processo di vendita Ã¨ revisionato periodicamente?",
            "Il team segue costantemente il processo definito?"
        ],
        "controllo": "Esiste un manuale o risorsa scritta sul processo di vendita?"
    },
    "Script e Protocolli di Vendita": {
        "principali": [
            "Esistono script o linee guida per interagire con i clienti?",
            "Gli script vengono aggiornati col feedback del mercato?",
            "I venditori ricevono formazione sull'uso degli script?",
            "Gli script includono strumenti per gestire obiezioni?",
            "Ci sono protocolli di vendita standardizzati?"
        ],
        "controllo": "Avete anche esempi/casi studio o modelli email standardizzati?"
    },
    "Tecnologia e CRM": {
        "principali": [
            "C'Ã¨ un CRM per tracciare le interazioni con i clienti?",
            "Il team aggiorna costantemente il CRM?",
            "Ci sono strumenti per automatizzare alcune fasi di vendita?",
            "Il CRM offre report utili per decisioni strategiche?",
            "Il team Ã¨ formato e motivato a usare il CRM?"
        ],
        "controllo": "C'Ã¨ una figura di riferimento per la gestione del CRM?"
    },
    "Formazione e Sviluppo del Team di Vendita": {
        "principali": [
            "Sono previste sessioni di formazione periodiche?",
            "Esiste un programma di onboarding strutturato?",
            "Il team riceve coaching/mentoring personalizzato?",
            "C'Ã¨ un piano di sviluppo a lungo termine per i venditori?",
            "Con che frequenza formi i venditori? (1=mai,5=molto frequente)"
        ],
        "controllo": "Hai un calendario formativo scritto con date e argomenti?"
    },
    "Misurazione e KPI": {
        "principali": [
            "Sono definiti KPI chiari e comprensibili?",
            "I KPI sono comunicati e compresi dal team?",
            "Si analizzano regolarmente i risultati dei KPI?",
            "Ci sono report/dashboard per monitorare i KPI?",
            "Si intraprendono azioni correttive in base ai KPI?"
        ],
        "controllo": "Il team riceve regolarmente un report con KPI e commenti?"
    }
}

st.title("Assessment Reparto Commerciale")
st.write("Compila le seguenti domande per ogni area, poi clicca su 'Genera Grafici' per visualizzare i risultati.")

nome_cliente = st.text_input("Nome del Cliente:", "Francesco Ramundo")
nome_azienda = st.text_input("Nome dell'Azienda:", "rarosrl.com")

punteggi_aree = {}
controlli_aree = {}

for area in aree:
    st.subheader(area)
    punteggi_principali = []
    for i, domanda_principale in enumerate(domande[area]["principali"]):
        val = selettore_valori_domanda(domanda_principale, key=area+str(i))
        punteggi_principali.append(val)

    domanda_controllo = domande[area]["controllo"]
    val_controllo = selettore_valori_domanda("(Controllo) " + domanda_controllo, key=area+"controllo")

    punteggi_aree[area] = punteggi_principali
    controlli_aree[area] = val_controllo

def crea_radar(labels, values, title, max_score=5):
    vals_perc = [(v/max_score)*100 for v in values]
    N = len(labels)
    angles = np.linspace(0, 2*np.pi, N, endpoint=False)
    vals_perc += vals_perc[:1]
    angles = np.concatenate((angles, [angles[0]]))

    fig, ax = plt.subplots(figsize=(12,12), subplot_kw=dict(polar=True))
    ax.set_theta_offset(np.pi/2 + 0.1)
    ax.set_theta_direction(-1)

    ax.fill(angles, vals_perc, color='blue', alpha=0.25)
    ax.plot(angles, vals_perc, color='blue', linewidth=2)

    for angle in angles[:-1]:
        ax.plot([angle, angle],[0,100], color='gray', linewidth=1)

    ax.set_ylim(0,100)

    for valp, angle in zip(vals_perc, angles):
        ax.text(angle, valp+5, f'{valp:.0f}%', horizontalalignment='center', verticalalignment='bottom', size=9, color='black')

    for label, angle in zip(labels, angles[:-1]):
        ax.text(angle,120,label, horizontalalignment='center', verticalalignment='center', size=10, wrap=True)

    valore_ottimale = 85
    ottimale_stats = [valore_ottimale]*(N+1)
    ax.plot(angles, ottimale_stats, color='orange', linestyle='dashed', linewidth=2)

    plt.title(title, fontsize=16, pad=80)
    ax.set_xticklabels([])
    ax.set_thetagrids([])

    return fig

if st.button("Genera Grafici"):
    medie_aree = {}
    for area in aree:
        media = np.mean(punteggi_aree[area])
        controllo = controlli_aree[area]
        if media >=4 and controllo <=2:
            media = max(1, media - 1)
            st.warning(f"L'area '{area}' aveva media {np.mean(punteggi_aree[area]):.2f}, ma a causa della bassa risposta di controllo ({controllo}), il punteggio Ã¨ stato ridotto a {media:.2f}.")
        medie_aree[area] = media

    st.header("Risultati Generali")
    fig_generale = crea_radar(aree, [medie_aree[a] for a in aree], "Analisi Reparto Commerciale")
    st.pyplot(fig_generale)

    st.header("Dettaglio per Area")
    aree_files = {}
    for area in aree:
        fig_area = crea_radar(domande[area]["principali"], punteggi_aree[area], area)
        st.pyplot(fig_area)
        aree_files[area] = fig_area

    base_dir = "Programma Test vendita"
    try:
        if not os.path.exists(base_dir):
            os.makedirs(base_dir)

        cliente_dir = os.path.join(base_dir, nome_cliente)
        if not os.path.exists(cliente_dir):
            os.makedirs(cliente_dir)
    except Exception as e:
        st.error(f"Errore nella creazione delle directory: {e}")
        st.stop()

    try:
        fig_generale.savefig(os.path.join(cliente_dir, "grafico_generale.png"), dpi=300, bbox_inches='tight')
        for area in aree:
            fig_area = aree_files[area]
            fig_area.savefig(os.path.join(cliente_dir, f"grafico_{area}.png"), dpi=300, bbox_inches='tight')
    except Exception as e:
        st.error(f"Errore nel salvataggio dei grafici: {e}")
        st.stop()

    doc = Document()

    # Percorso assoluto del logo (assicurarsi che esista)
    logo_path = "Logo-Sales-Flow-payoff.png"
    if not os.path.exists(logo_path):
        st.error(f"Il file del logo non esiste nel percorso: {logo_path}")
    else:
        try:
            section = doc.sections[0]
            header = section.header
            header_par = header.paragraphs[0]
            header_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = header_par.add_run()
            r.add_picture(logo_path, width=Inches(1.5))
        except Exception as e:
            st.error(f"Errore nell'aggiunta del logo al documento: {e}")

    doc.add_heading('SALES ASSESSMENT', level=1)

    p_cliente = doc.add_paragraph()
    p_cliente.add_run("Report Analitico del Test di Vendita di ").bold = False
    run_cliente = p_cliente.add_run(nome_cliente)
    run_cliente.bold = False

    p_azienda = doc.add_paragraph()
    p_azienda.add_run("Azienda: ").bold = True
    run_azienda = p_azienda.add_run(nome_azienda)
    run_azienda.bold = False

    doc.add_paragraph("Questo documento fornisce un'analisi dettagliata della struttura commerciale dell'azienda, evidenziando punti di forza e aree di miglioramento.")

    # Nota modificata
    p_nota = doc.add_paragraph()
    run_nota = p_nota.add_run("Nota: ")
    run_nota.bold = True
    run_nota.font.color.rgb = RGBColor(0xFF,0x00,0x00)
    p_nota.add_run("Per garantire che il lavoro di ricerca venditori che svolgerÃ  RecruitFlow funzioni e per garantire che il venditore trovato riesca ad effettuare il proprio lavoro in modo efficace Ã¨ essenziale che le seguenti aree chiave raggiungano almeno l'85%:")

    doc.add_paragraph("- Visione e Strategia di Vendita")
    doc.add_paragraph("- Processi di Vendita")
    doc.add_paragraph("- Script e Protocolli di Vendita")
    doc.add_paragraph("- Formazione e Sviluppo del Team di Vendita")

    footer = section.footer
    footer_par = footer.paragraphs[0]
    footer_par.text = "Sales Assessmentâ„¢ SalesFlow - salesflow.it"
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading('Analisi Reparto Commerciale', level=2)
    try:
        doc.add_picture(os.path.join(cliente_dir, "grafico_generale.png"), width=Inches(6))
    except Exception as e:
        st.error(f"Errore nell'aggiunta del grafico generale al documento: {e}")

    for area in aree:
        doc.add_page_break()
        doc.add_heading(area, level=2)
        try:
            doc.add_picture(os.path.join(cliente_dir, f"grafico_{area}.png"), width=Inches(6))
        except Exception as e:
            st.error(f"Errore nell'aggiunta del grafico per l'area '{area}' al documento: {e}")

    # Salva il documento Word
    doc_name = f"report_assessment_{nome_cliente}.docx"
    doc_path = os.path.join(cliente_dir, doc_name)
    try:
        doc.save(doc_path)
    except Exception as e:
        st.error(f"Errore nel salvataggio del documento Word: {e}")
        st.stop()

    # Leggi il file Word in modalitÃ  binaria per il download
    try:
        with open(doc_path, "rb") as file:
            word_file = file.read()
    except Exception as e:
        st.error(f"Errore nella lettura del documento Word: {e}")
        st.stop()

    # Aggiungi il pulsante di download
    st.download_button(
        label="ðŸ“¥ Scarica il Report Word",
        data=word_file,
        file_name=doc_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.success(f"Il report Ã¨ stato generato con successo! Puoi scaricarlo usando il pulsante sopra.")
